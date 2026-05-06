"""Exports notes to DOCX using python-docx."""

from pathlib import Path

from docx import Document
from docx.shared import Pt


def export_full_note(
    title: str,
    transcript: str,
    summary: dict,
    category_name: str,
    output_path: Path,
) -> Path:
    doc = Document()

    doc.add_heading(title, level=1)
    doc.add_heading(f"Summary — {category_name}", level=2)

    for key, value in summary.items():
        heading = key.replace("_", " ").title()
        doc.add_heading(heading, level=3)

        if isinstance(value, str):
            doc.add_paragraph(value)
        elif isinstance(value, list):
            if value and isinstance(value[0], dict):
                for item in value:
                    doc.add_paragraph(
                        ", ".join(str(v) for v in item.values()),
                        style="List Bullet",
                    )
            else:
                for item in value:
                    doc.add_paragraph(str(item), style="List Bullet")
        elif isinstance(value, (bool, int)):
            doc.add_paragraph(str(value))
        else:
            doc.add_paragraph(str(value))

    doc.add_heading("Transcript", level=2)
    doc.add_paragraph(transcript)

    doc.save(str(output_path))
    return output_path


def export_transcript(transcript: str, output_path: Path) -> Path:
    doc = Document()
    doc.add_heading("Transcript", level=1)
    doc.add_paragraph(transcript)
    doc.save(str(output_path))
    return output_path
